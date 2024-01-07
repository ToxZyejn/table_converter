from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
import os


def set_cell_border(cell, **kwargs):
    """
    Установка границ ячейки таблицы
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Создаем элемент 'w:tcBorders' для границ
    tcBorders = OxmlElement('w:tcBorders')
    for key, value in kwargs.items():
        tag = 'w:{}'.format(key)
        element = OxmlElement(tag)
        element.set(qn('w:val'), value)
        tcBorders.append(element)

    tcPr.append(tcBorders)

def apply_borders_to_table(table):
    """
    Применяем границы ко всем ячейкам в таблице
    """
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top="single", bottom="single", start="single", end="single")

def extract_tables(source_filenames, keyword):
    new_doc = Document()
    for filename in source_filenames:
        doc = Document(filename)
        for table in doc.tables:
            if keyword in [cell.text for row in table.rows for cell in row.cells]:
                new_tbl = new_doc.add_table(rows=0, cols=len(table.columns))
                for row in table.rows:
                    cells = row.cells
                    new_row = new_tbl.add_row().cells
                    for idx, cell in enumerate(cells):
                        new_row[idx].text = cell.text
                apply_borders_to_table(new_tbl)
                new_doc.add_paragraph()

    new_doc.save('extracted_tables.docx')


# Список файлов для обработки
source_filenames = ['your_word_file_1.docx', ..., 'your_word_file_10.docx']

# Ключевое слово для поиска таблиц
keyword = 'keyword_for_searching_table'

# Вызов функции
extract_tables(source_filenames, keyword)


def select_files():
    global source_filenames
    filetypes = [('Word files', '*.docx'), ('All files', '*.*')]
    filenames = filedialog.askopenfilenames(title='Open files', initialdir='/', filetypes=filetypes)
    source_filenames = list(filenames)
    if source_filenames:
        files_label.config(text="\n".join(source_filenames))
    else:
        files_label.config(text="No files selected")

def start_extraction():
    keyword = simpledialog.askstring("Input", "Enter the keyword for searching tables",
                                     parent=root)
    if keyword and source_filenames:
        try:
            extract_tables(source_filenames, keyword)
            messagebox.showinfo("Success", "Tables extracted successfully!")
        except Exception as e:
            messagebox.showerror("Error", str(e))
    elif not source_filenames:
        messagebox.showwarning("Warning", "No files selected!")
    elif not keyword:
        messagebox.showwarning("Warning", "Keyword not entered!")

# GUI setup
root = tk.Tk()
root.title("Word Table Extractor")

select_button = tk.Button(root, text="Select Word Files", command=select_files)
select_button.pack()

files_label = tk.Label(root, text="No files selected")
files_label.pack()

start_button = tk.Button(root, text="Start Extraction", command=start_extraction)
start_button.pack()

root.mainloop()
